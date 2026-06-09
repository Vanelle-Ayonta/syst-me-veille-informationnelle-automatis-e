"""
core/export.py — Génération des rapports Word et PDF
Bulletin de veille institutionnel DIIF/BEAC
"""
import os, sys, io, logging
from datetime import datetime, timedelta
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from config import ANTHROPIC_API_KEY, get_model, DIMENSIONS_IF
from core.database import get_db

log = logging.getLogger(__name__)


def get_articles_periode(jours: int = 7) -> dict:
    """
    Récupère les articles de la période par dimension.
    Retourne un dict {dimension: [articles]}
    """
    depuis = (datetime.utcnow() - timedelta(days=jours)).isoformat()

    with get_db() as conn:
        articles = conn.execute("""
            SELECT a.id, a.titre, a.contenu, a.resume,
                   a.url_original, a.publie_le, a.collecte_le,
                   a.langue, a.qualite_contenu,
                   s.nom as source
            FROM articles a
            JOIN sources s ON a.source_id = s.id
            WHERE a.collecte_le >= ?
              AND a.contenu IS NOT NULL
              AND LENGTH(a.contenu) > 100
            ORDER BY a.collecte_le DESC
        """, (depuis,)).fetchall()

        sources = conn.execute("""
            SELECT s.nom, s.langue, s.url,
                   COUNT(a.id) as nb,
                   MAX(a.collecte_le) as derniere
            FROM sources s
            LEFT JOIN articles a ON a.source_id = s.id
              AND a.collecte_le >= ?
            GROUP BY s.id
            HAVING nb > 0
            ORDER BY nb DESC
        """, (depuis,)).fetchall()

    articles = [dict(r) for r in articles]
    sources  = [dict(r) for r in sources]

    # Classer par dimension
    par_dimension = {dim: [] for dim in DIMENSIONS_IF}
    par_dimension["Général"] = []

    mots_cles_par_dim = {
        "Accès":                        ["accès", "access", "account", "bancarisation", "unbanked"],
        "Utilisation":                  ["utilisation", "usage", "transaction", "paiement", "payment", "mobile money"],
        "Qualité":                      ["qualité", "quality", "fiabilité", "reliability", "transparence"],
        "Éducation financière":         ["éducation financière", "financial literacy", "sensibilisation", "alphabétisation"],
        "Protection des consommateurs": ["protection", "consumer", "fraude", "fraud", "sécurité", "security"],
        "Innovation financière":        ["innovation", "fintech", "digital", "numérique", "cbdc", "blockchain"],
    }

    for art in articles:
        texte = ((art.get("titre") or "") + " " +
                 (art.get("contenu") or "")).lower()
        classe = False
        for dim in DIMENSIONS_IF:
            mots = mots_cles_par_dim.get(dim, [dim.lower()])
            if any(m in texte for m in mots):
                par_dimension[dim].append(art)
                classe = True
                break
        if not classe:
            par_dimension["Général"].append(art)

    return {
        "articles":      articles,
        "par_dimension": par_dimension,
        "sources":       sources,
        "total":         len(articles),
        "depuis":        depuis[:10],
        "jusqu":         datetime.utcnow().strftime("%Y-%m-%d"),
    }


def generer_synthese_editoriale(
    articles: list,
    jours: int,
    suggestions: list = None
) -> str:
    """
    Génère la synthèse éditoriale via Claude.
    Style journal institutionnel.
    """
    if not articles:
        return (
            "Aucun article collecté sur cette période. "
            "Veuillez lancer la collecte."
        )

    import anthropic
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

    extraits = "\n\n".join([
        f"[{art['source']}] {art['titre']}\n"
        f"{(art.get('resume') or art.get('contenu','')[:300])}"
        for art in articles[:15]
    ])

    prompt = f"""Tu es rédacteur en chef du bulletin de veille du DIIF
(Département de l'Inclusion et de l'Innovation Financières) de la BEAC.

Rédige une synthèse éditoriale professionnelle et fluide des actualités
de l'inclusion financière des {jours} derniers jours.

Style : journal institutionnel, paragraphes narratifs, ton factuel
et analytique, environ 400 mots.

Structure obligatoire :
1. Ouverture : fait marquant principal de la période
2. Tendances : 2-3 tendances clés identifiées
3. Zone CEMAC : focus sur l'Afrique centrale si présent
4. Perspectives : ce qu'il faut surveiller

Articles de la période :
{extraits}

Rédige uniquement la synthèse, sans titre ni en-tête."""

    try:
        response = client.messages.create(
            model=get_model("smart"),
            max_tokens=1500,
            messages=[{"role": "user", "content": prompt}],
        )
        return response.content[0].text.strip()
    except Exception as e:
        log.error(f"[EXPORT] Erreur synthèse : {e}")
        return (
            f"Synthèse de la période du "
            f"{datetime.utcnow().strftime('%d/%m/%Y')}.\n\n"
            f"{len(articles)} articles collectés depuis "
            f"{jours} sources."
        )


def generer_recommandations(
    articles: list,
    dimension: str = None
) -> list:
    """Génère les recommandations via l'agent SVIA."""
    try:
        from agents.svia_agent import orchestrer
        question = (
            "Quelles sont les recommandations prioritaires "
            "pour le DIIF/BEAC basées sur l'actualité récente "
            "de l'inclusion financière ?"
        )
        result = orchestrer(question=question, dimension=dimension, top_k=8)
        return result.get("suggestions", [])
    except Exception as e:
        log.warning(f"[EXPORT] Recommandations : {e}")
        return []


# ── Export Word ───────────────────────────────────────────────────

def generer_rapport_word(
    jours: int = 7,
    avec_recommandations: bool = True,
) -> bytes:
    """
    Génère le rapport Word institutionnel DIIF/BEAC.
    Retourne les bytes du fichier .docx.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    NAVY = RGBColor(0x00, 0x2B, 0x5C)
    GOLD = RGBColor(0xC8, 0xA9, 0x51)

    def add_heading(text, level=1, color=NAVY):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = color
            run.font.bold = True
        return p

    def add_paragraph(text, bold=False, italic=False, size=11):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size   = Pt(size)
        run.font.bold   = bold
        run.font.italic = italic
        return p

    def add_separator():
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "C8A951")
        pBdr.append(bottom)
        pPr.append(pBdr)

    donnees = get_articles_periode(jours)
    synthese = generer_synthese_editoriale(donnees["articles"], jours)
    recommandations = (
        generer_recommandations(donnees["articles"])
        if avec_recommandations else []
    )

    date_rapport = datetime.utcnow().strftime("%d/%m/%Y")
    periode = f"du {donnees['depuis']} au {donnees['jusqu']}"

    # ── PAGE DE GARDE ──────────────────────────────────────────────
    p_titre = doc.add_paragraph()
    p_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_titre.add_run("BANQUE DES ÉTATS DE L'AFRIQUE CENTRALE")
    run.font.size      = Pt(14)
    run.font.bold      = True
    run.font.color.rgb = NAVY

    p_dept = doc.add_paragraph()
    p_dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p_dept.add_run(
        "Département de l'Inclusion et de l'Innovation Financières"
    )
    run2.font.size      = Pt(12)
    run2.font.color.rgb = GOLD

    doc.add_paragraph()
    add_separator()
    doc.add_paragraph()

    p_bull = doc.add_paragraph()
    p_bull.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p_bull.add_run("BULLETIN DE VEILLE INFORMATIONNELLE")
    run3.font.size      = Pt(18)
    run3.font.bold      = True
    run3.font.color.rgb = NAVY

    p_per = doc.add_paragraph()
    p_per.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p_per.add_run(f"Période : {periode}")
    run4.font.size      = Pt(12)
    run4.font.color.rgb = GOLD

    doc.add_paragraph()
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p_date.add_run(f"Généré le {date_rapport}")
    run5.font.size      = Pt(10)
    run5.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    p_total = doc.add_paragraph()
    p_total.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run6 = p_total.add_run(
        f"{donnees['total']} articles analysés · "
        f"{len(donnees['sources'])} sources"
    )
    run6.font.size   = Pt(10)
    run6.font.italic = True

    doc.add_page_break()

    # ── SECTION 1 : SYNTHÈSE ÉDITORIALE ───────────────────────────
    add_heading("1. Synthèse de la période", level=1)
    add_separator()
    doc.add_paragraph()

    for para in synthese.split("\n\n"):
        if para.strip():
            add_paragraph(para.strip())
            doc.add_paragraph()

    doc.add_page_break()

    # ── SECTION 2 : ACTUALITÉS PAR DIMENSION ──────────────────────
    add_heading("2. Actualités par dimension", level=1)
    add_separator()

    all_dims = DIMENSIONS_IF + ["Général"]
    for idx, dim in enumerate(all_dims):
        arts = donnees["par_dimension"].get(dim, [])
        if not arts:
            continue
        doc.add_paragraph()
        num = idx + 1 if dim in DIMENSIONS_IF else len(DIMENSIONS_IF) + 1
        add_heading(f"2.{num}. {dim}", level=2)

        for art in arts[:5]:
            p = doc.add_paragraph(style="List Bullet")
            run_titre = p.add_run(art["titre"] or "Sans titre")
            run_titre.font.bold = True
            run_titre.font.size = Pt(10)

            p2 = doc.add_paragraph()
            run_meta = p2.add_run(
                f"Source : {art['source']} · "
                f"Date : {(art.get('publie_le') or art.get('collecte_le',''))[:10]}"
            )
            run_meta.font.size      = Pt(9)
            run_meta.font.italic    = True
            run_meta.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

            resume = art.get("resume") or (art.get("contenu") or "")[:300]
            if resume:
                p3 = doc.add_paragraph()
                run_res = p3.add_run(resume[:400])
                run_res.font.size = Pt(10)

            if art.get("url_original"):
                p4 = doc.add_paragraph()
                run_url = p4.add_run(f"Lien : {art['url_original']}")
                run_url.font.size      = Pt(9)
                run_url.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)

            doc.add_paragraph()

    doc.add_page_break()

    # ── SECTION 3 : SOURCES CONSULTÉES ────────────────────────────
    add_heading("3. Sources consultées", level=1)
    add_separator()
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    headers = ["Source", "Langue", "Articles", "Lien"]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold      = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc   = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "002B5C")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)

    for src in donnees["sources"]:
        row = table.add_row().cells
        row[0].text = src["nom"]
        row[1].text = src["langue"].upper()
        row[2].text = str(src["nb"])
        row[3].text = src.get("url", "")

    doc.add_page_break()

    # ── SECTION 4 : RECOMMANDATIONS ───────────────────────────────
    if recommandations:
        add_heading("4. Recommandations DIIF", level=1)
        add_separator()
        doc.add_paragraph()

        priorite_label = {
            "haute":   "PRIORITÉ HAUTE",
            "moyenne": "PRIORITÉ MOYENNE",
            "faible":  "PRIORITÉ FAIBLE",
        }

        for i, rec in enumerate(recommandations, 1):
            p_rec = doc.add_paragraph()
            run_rec = p_rec.add_run(f"{i}. {rec.get('titre', '')}")
            run_rec.font.bold      = True
            run_rec.font.size      = Pt(11)
            run_rec.font.color.rgb = NAVY

            prio = rec.get("priorite", "moyenne")
            p_prio = doc.add_paragraph()
            run_prio = p_prio.add_run(
                f"{priorite_label.get(prio, prio.upper())} · "
                f"{rec.get('dimension', '')}"
            )
            run_prio.font.size      = Pt(9)
            run_prio.font.italic    = True
            run_prio.font.color.rgb = GOLD

            p_desc = doc.add_paragraph()
            run_desc = p_desc.add_run(rec.get("description", ""))
            run_desc.font.size = Pt(10)
            doc.add_paragraph()

    # ── PIED DE PAGE ───────────────────────────────────────────────
    section = doc.sections[0]
    footer  = section.footer
    p_foot  = footer.paragraphs[0]
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_foot = p_foot.add_run(
        f"DIIF / BEAC — Document confidentiel — "
        f"Généré le {date_rapport}"
    )
    run_foot.font.size      = Pt(8)
    run_foot.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── Export PDF ────────────────────────────────────────────────────

def generer_rapport_pdf(jours: int = 7) -> bytes:
    """
    Génère le rapport PDF institutionnel DIIF/BEAC.
    Retourne les bytes du fichier .pdf.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor, white
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer,
        Table, TableStyle, PageBreak, HRFlowable
    )
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

    NAVY_RL = HexColor("#002B5C")
    GOLD_RL = HexColor("#C8A951")
    GREY_RL = HexColor("#888888")

    donnees = get_articles_periode(jours)
    synthese = generer_synthese_editoriale(donnees["articles"], jours)
    recommandations = generer_recommandations(donnees["articles"])

    date_rapport = datetime.utcnow().strftime("%d/%m/%Y")
    periode = f"du {donnees['depuis']} au {donnees['jusqu']}"

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=2*cm, bottomMargin=2*cm,
        leftMargin=2.5*cm, rightMargin=2.5*cm,
    )

    styles = getSampleStyleSheet()
    style_titre = ParagraphStyle(
        "titre", parent=styles["Title"],
        textColor=NAVY_RL, fontSize=18,
        spaceAfter=12, alignment=TA_CENTER,
    )
    style_h1 = ParagraphStyle(
        "h1", parent=styles["Heading1"],
        textColor=NAVY_RL, fontSize=14,
        spaceBefore=16, spaceAfter=8,
    )
    style_h2 = ParagraphStyle(
        "h2", parent=styles["Heading2"],
        textColor=NAVY_RL, fontSize=12,
        spaceBefore=12, spaceAfter=6,
    )
    style_body = ParagraphStyle(
        "body", parent=styles["Normal"],
        fontSize=10, leading=14,
        spaceAfter=6, alignment=TA_JUSTIFY,
    )
    style_meta = ParagraphStyle(
        "meta", parent=styles["Normal"],
        fontSize=8, textColor=GREY_RL,
        spaceAfter=4,
    )
    style_center = ParagraphStyle(
        "center", parent=styles["Normal"],
        fontSize=11, alignment=TA_CENTER,
        textColor=NAVY_RL,
    )
    style_gold = ParagraphStyle(
        "gold", parent=styles["Normal"],
        fontSize=12, alignment=TA_CENTER,
        textColor=GOLD_RL,
    )

    story = []

    # Page de garde
    story.append(Spacer(1, 2*cm))
    story.append(Paragraph(
        "BANQUE DES ÉTATS DE L'AFRIQUE CENTRALE", style_center
    ))
    story.append(Paragraph(
        "Département de l'Inclusion et de l'Innovation Financières",
        style_gold
    ))
    story.append(Spacer(1, 0.5*cm))
    story.append(HRFlowable(
        width="100%", thickness=2, color=GOLD_RL, spaceAfter=12
    ))
    story.append(Spacer(1, 1*cm))
    story.append(Paragraph("BULLETIN DE VEILLE INFORMATIONNELLE", style_titre))
    story.append(Paragraph(f"Période : {periode}", style_gold))
    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph(
        f"Généré le {date_rapport} · "
        f"{donnees['total']} articles · "
        f"{len(donnees['sources'])} sources",
        style_meta
    ))
    story.append(PageBreak())

    # Section 1 : Synthèse
    story.append(Paragraph("1. Synthèse de la période", style_h1))
    story.append(HRFlowable(
        width="100%", thickness=1, color=GOLD_RL, spaceAfter=8
    ))
    for para in synthese.split("\n\n"):
        if para.strip():
            story.append(Paragraph(para.strip(), style_body))
            story.append(Spacer(1, 6))
    story.append(PageBreak())

    # Section 2 : Actualités par dimension
    story.append(Paragraph("2. Actualités par dimension", style_h1))
    story.append(HRFlowable(
        width="100%", thickness=1, color=GOLD_RL, spaceAfter=8
    ))
    all_dims = DIMENSIONS_IF + ["Général"]
    for dim in all_dims:
        arts = donnees["par_dimension"].get(dim, [])
        if not arts:
            continue
        story.append(Paragraph(dim, style_h2))
        for art in arts[:4]:
            titre = art.get("titre") or "Sans titre"
            story.append(Paragraph(f"<b>{titre}</b>", style_body))
            date_art = (
                art.get("publie_le") or art.get("collecte_le", "")
            )[:10]
            story.append(Paragraph(
                f"{art['source']} · {date_art}", style_meta
            ))
            resume = art.get("resume") or (art.get("contenu") or "")[:250]
            if resume:
                story.append(Paragraph(resume[:350], style_body))
            if art.get("url_original"):
                story.append(Paragraph(
                    f"<link href='{art['url_original']}'>"
                    f"<font color='#0070C0'>"
                    f"{art['url_original'][:70]}"
                    f"</font></link>",
                    style_meta
                ))
            story.append(Spacer(1, 8))
    story.append(PageBreak())

    # Section 3 : Sources
    story.append(Paragraph("3. Sources consultées", style_h1))
    story.append(HRFlowable(
        width="100%", thickness=1, color=GOLD_RL, spaceAfter=8
    ))
    data_table = [["Source", "Langue", "Articles", "URL"]]
    for src in donnees["sources"]:
        data_table.append([
            src["nom"],
            src["langue"].upper(),
            str(src["nb"]),
            (src.get("url") or "")[:50],
        ])
    t = Table(data_table, colWidths=[4.5*cm, 2*cm, 2*cm, 8*cm])
    t.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1, 0), NAVY_RL),
        ("TEXTCOLOR",     (0, 0), (-1, 0), white),
        ("FONTNAME",      (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE",      (0, 0), (-1, -1), 8),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [HexColor("#F5F5F5"), white]),
        ("GRID",          (0, 0), (-1, -1), 0.5, GREY_RL),
        ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING",    (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(t)
    story.append(PageBreak())

    # Section 4 : Recommandations
    if recommandations:
        story.append(Paragraph("4. Recommandations DIIF", style_h1))
        story.append(HRFlowable(
            width="100%", thickness=1, color=GOLD_RL, spaceAfter=8
        ))
        for i, rec in enumerate(recommandations, 1):
            story.append(Paragraph(
                f"<b>{i}. {rec.get('titre','')}</b>", style_body
            ))
            prio    = rec.get("priorite", "moyenne").upper()
            dim_rec = rec.get("dimension", "")
            story.append(Paragraph(
                f"<font color='#C8A951'>Priorité {prio}"
                f"{' · ' + dim_rec if dim_rec else ''}</font>",
                style_meta
            ))
            story.append(Paragraph(rec.get("description", ""), style_body))
            story.append(Spacer(1, 8))

    doc.build(story)
    buf.seek(0)
    return buf.getvalue()


# ── Envoi email bulletin ──────────────────────────────────────────

def envoyer_bulletin_email(
    jours: int = 7,
    destinataires_extra: list = None
) -> dict:
    """
    Envoie le bulletin hebdomadaire par email.
    Destinataires : admins en base + liste extra.
    """
    import smtplib
    from email.mime.text import MIMEText
    from email.mime.multipart import MIMEMultipart
    from email.mime.base import MIMEBase
    from email import encoders
    from config import (SMTP_HOST, SMTP_PORT,
                        SMTP_USER, SMTP_PASSWORD, APP_URL)

    with get_db() as conn:
        admins = conn.execute("""
            SELECT email, nom FROM utilisateurs
            WHERE role = 'administrateur'
              AND email IS NOT NULL
              AND email != ''
        """).fetchall()

    destinataires = [r["email"] for r in admins]

    extra_env = os.getenv("MAIL_EXTRA_DESTINATAIRES", "")
    if extra_env:
        destinataires += [e.strip() for e in extra_env.split(",") if e.strip()]
    if destinataires_extra:
        destinataires += destinataires_extra

    destinataires = list(set(destinataires))

    if not destinataires:
        return {"success": False, "error": "Aucun destinataire configuré."}

    donnees         = get_articles_periode(jours)
    synthese        = generer_synthese_editoriale(donnees["articles"], jours)
    recommandations = generer_recommandations(donnees["articles"])

    date_rapport = datetime.utcnow().strftime("%d/%m/%Y")
    periode      = f"du {donnees['depuis']} au {donnees['jusqu']}"

    recs_html = ""
    for rec in recommandations[:3]:
        prio_couleur = {
            "haute":   "#E74C3C",
            "moyenne": "#E67E22",
            "faible":  "#27AE60",
        }.get(rec.get("priorite", ""), "#888")
        recs_html += f"""
        <div style="border-left:3px solid {prio_couleur};
                    padding:8px 12px;margin-bottom:8px;
                    background:#fafafa;">
            <strong>{rec.get('titre','')}</strong><br>
            <small style="color:{prio_couleur};">
                Priorité {rec.get('priorite','').upper()}
            </small><br>
            <span style="font-size:13px;">
                {rec.get('description','')[:200]}
            </span>
        </div>"""

    synthese_court = "\n\n".join(synthese.split("\n\n")[:2])

    corps_html = f"""
    <html><body style="font-family:Arial,sans-serif;
                       color:#333;max-width:650px;margin:0 auto;">
        <div style="background:#002B5C;padding:1.5rem 2rem;
                    border-bottom:3px solid #C8A951;">
            <h1 style="color:#fff;margin:0;font-size:20px;">
                Bulletin de Veille DIIF/BEAC
            </h1>
            <p style="color:#C8A951;margin:4px 0 0;">
                Période : {periode}
            </p>
        </div>
        <div style="padding:1.5rem 2rem;background:#f9f9f9;">
            <p style="color:#888;font-size:12px;">
                Généré le {date_rapport} ·
                {donnees['total']} articles ·
                {len(donnees['sources'])} sources
            </p>
            <h2 style="color:#002B5C;
                       border-bottom:2px solid #C8A951;
                       padding-bottom:6px;">
                Synthèse de la période
            </h2>
            <p style="line-height:1.6;">
                {synthese_court.replace(chr(10), '<br>')}
            </p>
            <p>
                <a href="{APP_URL}"
                   style="background:#002B5C;color:#fff;
                          padding:8px 16px;border-radius:6px;
                          text-decoration:none;font-weight:600;">
                    Lire le rapport complet →
                </a>
            </p>
            <h2 style="color:#002B5C;
                       border-bottom:2px solid #C8A951;
                       padding-bottom:6px;">
                Recommandations prioritaires
            </h2>
            {recs_html if recs_html else "<p>Aucune recommandation générée.</p>"}
        </div>
        <div style="padding:1rem 2rem;background:#002B5C;text-align:center;">
            <p style="color:rgba(255,255,255,0.6);font-size:11px;margin:0;">
                DIIF / BEAC — Système de veille automatisée<br>
                Ce message est confidentiel.
            </p>
        </div>
    </body></html>"""

    try:
        word_bytes  = generer_rapport_word(jours)
        nom_fichier = (
            f"Bulletin_Veille_DIIF_"
            f"{datetime.utcnow().strftime('%Y%m%d')}.docx"
        )
    except Exception as e:
        log.warning(f"[EMAIL] Word non généré : {e}")
        word_bytes  = None
        nom_fichier = None

    envoyes = []
    erreurs = []

    for dest in destinataires:
        try:
            msg = MIMEMultipart("mixed")
            msg["Subject"] = f"Bulletin de veille DIIF/BEAC — {periode}"
            msg["From"]    = f"Veille DIIF/BEAC <{SMTP_USER}>"
            msg["To"]      = dest

            msg.attach(MIMEText(corps_html, "html", "utf-8"))

            if word_bytes:
                part = MIMEBase(
                    "application",
                    "vnd.openxmlformats-officedocument"
                    ".wordprocessingml.document"
                )
                part.set_payload(word_bytes)
                encoders.encode_base64(part)
                part.add_header(
                    "Content-Disposition",
                    f'attachment; filename="{nom_fichier}"'
                )
                msg.attach(part)

            with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as srv:
                srv.ehlo()
                srv.starttls()
                srv.login(SMTP_USER, SMTP_PASSWORD)
                srv.sendmail(SMTP_USER, dest, msg.as_string())

            envoyes.append(dest)

        except Exception as e:
            log.error(f"[EMAIL] Erreur → {dest} : {e}")
            erreurs.append({"email": dest, "erreur": str(e)})

    return {
        "success": len(envoyes) > 0,
        "envoyes": envoyes,
        "erreurs": erreurs,
        "total":   len(destinataires),
    }